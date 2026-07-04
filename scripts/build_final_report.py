from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Informe_Final_Autonomous_Software_Factory_Juan_Jose_Cordova.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def add_code_box(doc, title, text):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    p.add_run(title)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F2F4F7")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Proyecto Final: Autonomous Software Factory")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Ingenieria de Software - Maestria en Software").bold = True
    subtitle.add_run("\nJuan Jose Cordova\nFecha de entrega: 5 de julio de 2026\nModalidad: individual")

    doc.add_paragraph()
    add_table(
        doc,
        ["Elemento", "Detalle"],
        [
            ["Producto trabajado", "Sistema de reserva de citas dentales"],
            ["Flujo integrado", "Discovery -> Delivery -> SDD -> Quality Gate"],
            ["Funcionalidad nueva", "US-06 Anti doble-agendamiento"],
            ["Stack", "Java 21, Spring Boot, Gradle, JaCoCo, Semgrep"],
            ["Resultado final", "Gate aprobado con 7/7 pruebas y cobertura 90.9%"],
        ],
    )

    doc.add_heading("Introduccion", level=1)
    doc.add_paragraph(
        "Este informe integra los agentes trabajados durante la materia en un flujo completo de "
        "Autonomous Software Factory. El caso usado fue un sistema de reserva de citas dentales. "
        "La idea se descubrio con el Discovery Agent, se organizo con el Agile Delivery Team, se "
        "implemento una funcionalidad mediante SDD y finalmente se valido con el Quality Agent."
    )

    doc.add_heading("Analisis del Discovery Agent", level=1)
    doc.add_paragraph(
        "En la Unidad 1 utilice entrevistas de doctor, paciente y secretaria para descubrir el problema "
        "principal del producto. El hallazgo central fue que la agenda estaba dispersa entre WhatsApp, "
        "llamadas, papel y atencion presencial, sin una fuente unica de verdad."
    )
    add_bullets(
        doc,
        [
            "Personas identificadas: doctor, paciente y secretaria.",
            "Dolores principales: doble agendamiento, no-shows, falta de confirmacion y poca visibilidad de horarios.",
            "MVP definido: reserva autonoma de citas dentales sobre horarios reales.",
            "Metrica principal: reducir la tasa de no-show y aumentar citas gestionadas dentro del sistema.",
        ],
    )
    add_code_box(
        doc,
        "Prompt utilizado",
        "Usar las entrevistas de doctor, paciente y secretaria para descubrir el MVP de un sistema de reserva de citas dentales. Identificar personas, dolores, propuesta de valor, MVP Canvas, riesgos, supuestos e historias iniciales.",
    )
    add_code_box(
        doc,
        "Log / evidencia",
        "Repositorio: https://github.com/juancvxdev/discovery-agent\nArtefactos: discoveries/citasdentista/outputs/mvp-canvas.md, personas.md, user-stories.md, report.html\nResultado: MVP Canvas de citasdentista con propuesta de valor y metrica de no-show.",
    )

    doc.add_heading("Analisis del Agile Delivery Team", level=1)
    doc.add_paragraph(
        "En la Unidad 2 use los resultados del Discovery Agent como insumo para organizar el MVP en epicas, "
        "historias INVEST y un plan de sprint. El trabajo separo el valor del producto en reserva autonoma, "
        "confirmacion, fuente unica de recepcion y agenda del doctor."
    )
    add_table(
        doc,
        ["Epica", "Resultado"],
        [
            ["E-01", "Paciente reserva su cita viendo turnos reales"],
            ["E-02", "Paciente confirma, recuerda y libera su cita"],
            ["E-03", "Recepcion opera sobre una fuente unica, sin choques"],
            ["E-04", "Doctor consulta agenda en tiempo real con contexto"],
        ],
    )
    doc.add_paragraph(
        "Para el trabajo nuevo seleccione la historia US-06 Anti doble-agendamiento, porque representa un "
        "riesgo real de operacion y obliga a validar concurrencia, no solo el camino feliz."
    )
    add_code_box(
        doc,
        "Prompt utilizado",
        "Tomar el MVP Canvas, personas, requisitos e historias iniciales del Discovery Agent para generar epicas, historias de usuario INVEST, criterios de aceptacion, backlog priorizado y plan de sprint.",
    )
    add_code_box(
        doc,
        "Log / evidencia",
        "Repositorio: https://github.com/juancvxdev/agile-delivery-team\nArtefactos: deliveries/citasdentista/outputs/epics.md, stories.md, sprint-plan.md, architecture.md\nHistoria seleccionada: US-06 Anti doble-agendamiento.",
    )

    doc.add_heading("Analisis del desarrollo con SDD", level=1)
    doc.add_paragraph(
        "La funcionalidad US-06 se implemento con Spec-Driven Development. Primero defini los criterios en "
        "Spec-Kit y luego implemente el servicio Spring Boot. Los artefactos quedaron versionados en "
        "`specs/us-06-anti-doble-agendamiento/`."
    )
    add_table(
        doc,
        ["Artefacto", "Proposito"],
        [
            ["spec.md", "Define FR-001 a FR-007, escenarios y edge cases"],
            ["plan.md", "Explica stack, diseno tecnico y validacion"],
            ["tasks.md", "Lista tareas de implementacion, pruebas y calidad"],
        ],
    )
    doc.add_paragraph(
        "El diseno usa `ConcurrentHashMap.putIfAbsent` para que el turno dental se reserve de forma atomica. "
        "Asi, cuando dos pacientes intentan tomar el mismo horario, solo una reserva queda registrada."
    )
    add_code_box(
        doc,
        "Prompt utilizado",
        "Implementar la historia US-06 Anti doble-agendamiento usando Spec-Driven Development con Spec-Kit. Crear spec.md, plan.md y tasks.md; luego implementar Java + Spring Boot con pruebas automatizadas.",
    )
    add_code_box(
        doc,
        "Log / evidencia",
        "Pruebas ejecutadas: 7\nFallos: 0\nErrores: 0\nCobertura JaCoCo: 90.9%\nArchivos principales: AgendaDentalService.java, AgendaDentalController.java, AgendaDentalServiceTest.java",
    )

    doc.add_heading("Analisis del Quality Agent", level=1)
    doc.add_paragraph(
        "La validacion se hizo con el Quality Agent de la Unidad 3, reutilizando su constitucion, skill, "
        "subagentes, comandos, hook y conexion MCP. El agente valido los tres pilares del Definition of Done: "
        "pruebas, seguridad y criterios."
    )
    add_table(
        doc,
        ["Pilar", "Resultado final"],
        [
            ["Pruebas", "7/7 pruebas aprobadas, cobertura 90.9%"],
            ["Seguridad", "Semgrep sin hallazgos, 0 criticas y 0 secretos"],
            ["Criterios", "FR-001 a FR-007 y edge case de concurrencia cubiertos"],
        ],
    )
    add_code_box(
        doc,
        "Prompt utilizado",
        "Validar el proyecto Spring Boot de citas dentales con el Quality Agent. Leer criterios desde specs/us-06-anti-doble-agendamiento/spec.md, ejecutar pruebas/cobertura, revisar seguridad y generar verification.json y report.html.",
    )
    add_code_box(
        doc,
        "Gate bloqueado",
        "GATE DE CALIDAD: BLOQUEADO\nPRUEBAS: OK\nSEGURIDAD: OK\nCRITERIOS: FR-003 incumple; EDGE-CONCURRENCY incumple\nMotivo: faltaba prueba explicita de concurrencia.",
    )
    add_code_box(
        doc,
        "Gate aprobado",
        "GATE DE CALIDAD: APROBADO\nPRUEBAS: 7/7 - cobertura 90.9% >= 80%\nSEGURIDAD: 0 criticas - 0 secretos\nCRITERIOS: 8 criterios cumplen",
    )

    doc.add_heading("Conclusiones", level=1)
    add_bullets(
        doc,
        [
            "El flujo completo conecto descubrimiento, planificacion, desarrollo y validacion automatica.",
            "La historia US-06 demostro que la concurrencia debe probarse explicitamente y no asumirse por cobertura.",
            "El Quality Gate cambio el cierre del trabajo: la funcionalidad solo se considero terminada cuando los tres pilares estuvieron en verde.",
        ],
    )

    doc.add_heading("Recomendaciones", level=1)
    add_bullets(
        doc,
        [
            "Mantener Spec-Kit como fuente de criterios antes de implementar nuevas historias.",
            "Ejecutar el Quality Agent antes de cada entrega para evitar omisiones de pruebas o seguridad.",
            "Agregar el gate a integracion continua cuando el producto avance a mas funcionalidades.",
        ],
    )

    doc.add_heading("Anexos", level=1)
    add_table(
        doc,
        ["Repositorio", "Enlace"],
        [
            ["Discovery Agent", "https://github.com/juancvxdev/discovery-agent"],
            ["Agile Delivery Team", "https://github.com/juancvxdev/agile-delivery-team"],
            ["Quality Agent", "https://github.com/juancvxdev/quality-agent"],
            ["Proyecto final SDD + evidencia", "https://github.com/juancvxdev/Autonomous-Software-Factory"],
        ],
    )
    add_bullets(
        doc,
        [
            "Evidencia SDD: specs/us-06-anti-doble-agendamiento/spec.md, plan.md y tasks.md.",
            "Evidencia Quality Agent: quality-output/verification.json y quality-output/report.html.",
            "Evidencia del bloqueo: quality-output/before-fix/gate-output.txt.",
            "Evidencia de resolucion: quality-output/gate-output.txt.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
